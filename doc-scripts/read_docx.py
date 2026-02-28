from docx import Document
import sys

doc_path = '/root/books/腾讯云大数据处理套件 TBDS_5.3.1.3_用户指南_01.docx'

try:
    doc = Document(doc_path)
    
    # 获取文档基本信息
    print(f"文档段落数: {len(doc.paragraphs)}")
    print(f"文档表格数: {len(doc.tables)}")
    print("\n" + "="*60)
    print("文档内容预览 (前100段):")
    print("="*60 + "\n")
    
    # 读取前100段
    for i, para in enumerate(doc.paragraphs[:100]):
        if para.text.strip():
            print(f"{i+1}. {para.text}")
    
    print("\n" + "="*60)
    print("文档结构/标题:")
    print("="*60 + "\n")
    
    # 提取标题（通常是粗体或特定样式）
    headings = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and len(text) < 100:  # 可能是标题
            if para.style.name.startswith('Heading') or '标题' in para.style.name:
                headings.append(text)
            elif text.startswith('第') and ('章' in text or '节' in text):
                headings.append(text)
    
    for h in headings[:50]:  # 只显示前50个标题
        print(f"- {h}")
        
except Exception as e:
    print(f"错误: {e}")
    sys.exit(1)
