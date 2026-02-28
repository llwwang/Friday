from docx import Document

doc_path = '/root/books/腾讯云大数据处理套件 TBDS_5.3.1.3_用户指南_01.docx'
doc = Document(doc_path)

print("="*60)
print("TBDS 关键概念和操作流程:")
print("="*60 + "\n")

# 读取关键段落
important_sections = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text and len(text) > 20 and len(text) < 300:
        # 查找重要概念
        if any(keyword in text for keyword in ['标准集群', '虚拟集群', '公共集群', '组件', '服务', 'Hive', 'Spark', 'HDFS', 'YARN']):
            if text not in [s[1] for s in important_sections]:
                important_sections.append((i, text))

# 显示前40条
for idx, text in important_sections[:40]:
    print(f"• {text}")

print("\n" + "="*60)
print("集群组件列表:")
print("="*60 + "\n")

# 查找组件相关的内容
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith('Hive') or text.startswith('Spark') or text.startswith('HDFS') or text.startswith('YARN'):
        if len(text) < 200:
            print(f"- {text[:100]}")

print("\n" + "="*60)
print("用户和权限管理:")
print("="*60 + "\n")

# 查找用户权限相关内容
for i, para in enumerate(doc.paragraphs[300:800]):
    text = para.text.strip()
    if any(keyword in text for keyword in ['用户', '权限', '角色', '认证', 'kerberos', 'ranger']):
        if len(text) > 20 and len(text) < 200:
            print(f"• {text}")
