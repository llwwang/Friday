from docx import Document

doc_path = '/root/books/腾讯云大数据处理套件 TBDS_5.3.1.3_用户指南_01.docx'
doc = Document(doc_path)

print("="*60)
print("TBDS Manager 相关内容:")
print("="*60 + "\n")

# 查找包含 "TM"、"Manager"、"运维"、"监控" 等关键词的段落
keywords = ['TBDS Manager', 'TM', '运维中心', '监控', '告警', '集群管理', '快速入门', '操作步骤']
found_sections = []

for i, para in enumerate(doc.paragraphs[100:500]):  # 读取100-500段
    text = para.text.strip()
    if text:
        for keyword in keywords:
            if keyword in text and len(text) < 200:
                found_sections.append((i+100, text))
                break

# 去重并显示
seen = set()
for idx, text in found_sections[:30]:
    if text not in seen:
        seen.add(text)
        print(f"{idx}. {text}")

print("\n" + "="*60)
print("文档主要章节结构:")
print("="*60 + "\n")

# 提取所有章节标题
for para in doc.paragraphs:
    text = para.text.strip()
    if text and (text.startswith('第') or text.startswith('1.') or text.startswith('2.') or text.startswith('3.')):
        if len(text) < 100 and ('章' in text or '节' in text or '.' in text[:3]):
            print(text)
